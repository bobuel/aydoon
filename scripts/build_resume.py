"""Build Alex Aidun's ATS-friendly executive resume as a DOCX."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = RGBColor(0x14, 0x21, 0x3D)
BLUE = RGBColor(0x33, 0x5C, 0xFF)
MUTED = RGBColor(0x5C, 0x68, 0x78)
LINE = "DFE4DC"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT = "Arial"

# compact_reference_guide with one named override: executive_resume_density.
# The override keeps a senior, evidence-rich career history to two readable pages.
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_X_IN = 0.70
MARGIN_Y_IN = 0.62
CONTENT_WIDTH_IN = PAGE_WIDTH_IN - (2 * MARGIN_X_IN)
BODY_SIZE_PT = 9.4
BODY_LINE_SPACING = 1.05
BULLET_MARKER_IN = 0.18
BULLET_TEXT_IN = 0.38


def set_run_font(run, *, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url, *, color=BLUE, bold=False):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), FONT)
    run_fonts.set(qn("w:hAnsi"), FONT)
    run_properties.append(run_fonts)
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), str(color))
    run_properties.append(run_color)
    run_size = OxmlElement("w:sz")
    run_size.set(qn("w:val"), "18")
    run_properties.append(run_size)
    if bold:
        run_properties.append(OxmlElement("w:b"))
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph, field_name):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {field_name} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, display, end])
    set_run_font(run, size=8, color=MUTED)


def add_bottom_border(paragraph, color=LINE, size=8, space=3):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_custom_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(round(BULLET_TEXT_IN * 1440)))
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), str(round(BULLET_TEXT_IN * 1440)))
    indent.set(qn("w:hanging"), str(round((BULLET_TEXT_IN - BULLET_MARKER_IN) * 1440)))
    p_pr.extend([tabs, indent])
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "335CFF")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "16")
    r_pr.extend([color, size])
    level.extend([start, num_format, level_text, level_justification, p_pr, r_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_id_element = OxmlElement("w:abstractNumId")
    abstract_id_element.set(qn("w:val"), str(abstract_id))
    num.append(abstract_id_element)
    numbering.append(num)
    return num_id


def apply_bullet_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])
    p_pr.append(num_pr)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = BODY_LINE_SPACING

    heading = doc.styles["Heading 1"]
    heading.font.name = FONT
    heading._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    heading.font.size = Pt(10.3)
    heading.font.bold = True
    heading.font.color.rgb = BLUE
    heading.paragraph_format.space_before = Pt(7)
    heading.paragraph_format.space_after = Pt(3.5)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.keep_together = True

    name_style = doc.styles.add_style("Resume Name", WD_STYLE_TYPE.PARAGRAPH)
    name_style.font.name = FONT
    name_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    name_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    name_style.font.size = Pt(26)
    name_style.font.bold = True
    name_style.font.color.rgb = NAVY
    name_style.paragraph_format.space_after = Pt(1)

    title_style = doc.styles.add_style("Resume Title", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = FONT
    title_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    title_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    title_style.font.size = Pt(12.5)
    title_style.font.bold = True
    title_style.font.color.rgb = BLUE
    title_style.paragraph_format.space_after = Pt(4)

    contact_style = doc.styles.add_style("Resume Contact", WD_STYLE_TYPE.PARAGRAPH)
    contact_style.font.name = FONT
    contact_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    contact_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    contact_style.font.size = Pt(9)
    contact_style.font.color.rgb = MUTED
    contact_style.paragraph_format.space_after = Pt(7)

    profile_style = doc.styles.add_style("Executive Profile", WD_STYLE_TYPE.PARAGRAPH)
    profile_style.font.name = FONT
    profile_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    profile_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    profile_style.font.size = Pt(10)
    profile_style.font.color.rgb = NAVY
    profile_style.paragraph_format.line_spacing = 1.12
    profile_style.paragraph_format.space_after = Pt(1)

    compact_style = doc.styles.add_style("Compact Line", WD_STYLE_TYPE.PARAGRAPH)
    compact_style.font.name = FONT
    compact_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    compact_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    compact_style.font.size = Pt(9)
    compact_style.font.color.rgb = NAVY
    compact_style.paragraph_format.line_spacing = 1.05
    compact_style.paragraph_format.space_after = Pt(1)

    job_style = doc.styles.add_style("Job Header", WD_STYLE_TYPE.PARAGRAPH)
    job_style.font.name = FONT
    job_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    job_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    job_style.font.size = Pt(9.8)
    job_style.font.color.rgb = NAVY
    job_style.paragraph_format.space_before = Pt(4.5)
    job_style.paragraph_format.space_after = Pt(0)
    job_style.paragraph_format.keep_with_next = True
    job_style.paragraph_format.tab_stops.add_tab_stop(
        Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT
    )

    role_style = doc.styles.add_style("Role Line", WD_STYLE_TYPE.PARAGRAPH)
    role_style.font.name = FONT
    role_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    role_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    role_style.font.size = Pt(9)
    role_style.font.italic = True
    role_style.font.color.rgb = MUTED
    role_style.paragraph_format.space_after = Pt(1.5)
    role_style.paragraph_format.keep_with_next = True

    bullet_style = doc.styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.font.name = FONT
    bullet_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    bullet_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    bullet_style.font.size = Pt(BODY_SIZE_PT)
    bullet_style.font.color.rgb = NAVY
    bullet_style.paragraph_format.space_before = Pt(0)
    bullet_style.paragraph_format.space_after = Pt(2.1)
    bullet_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    bullet_style.paragraph_format.line_spacing = BODY_LINE_SPACING
    bullet_style.paragraph_format.keep_together = True


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph(style="Heading 1")
    run = paragraph.add_run(text.upper())
    set_run_font(run, size=10.3, color=BLUE, bold=True)
    add_bottom_border(paragraph, color=LINE, size=6, space=2)
    return paragraph


def add_bullet(doc, text, num_id):
    paragraph = doc.add_paragraph(style="Resume Bullet")
    apply_bullet_numbering(paragraph, num_id)
    run = paragraph.add_run(text)
    set_run_font(run, size=BODY_SIZE_PT, color=NAVY)
    return paragraph


def add_job(doc, *, company, period, role, location, bullets, num_id):
    heading = doc.add_paragraph(style="Job Header")
    company_run = heading.add_run(company.upper())
    set_run_font(company_run, size=9.8, color=NAVY, bold=True)
    heading.add_run("\t")
    period_run = heading.add_run(period)
    set_run_font(period_run, size=8.8, color=MUTED, bold=True)

    role_line = doc.add_paragraph(style="Role Line")
    role_run = role_line.add_run(role)
    set_run_font(role_run, size=9, color=MUTED, italic=True)
    if location:
        separator = role_line.add_run(f" | {location}")
        set_run_font(separator, size=9, color=MUTED)

    for bullet in bullets:
        add_bullet(doc, bullet, num_id)


def setup_page_furniture(doc):
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.left_margin = Inches(MARGIN_X_IN)
    section.right_margin = Inches(MARGIN_X_IN)
    section.top_margin = Inches(MARGIN_Y_IN)
    section.bottom_margin = Inches(MARGIN_Y_IN)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    section.different_first_page_header_footer = True

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.paragraph_format.space_after = Pt(0)
    name_run = header_paragraph.add_run("ALEXANDER AIDUN")
    set_run_font(name_run, size=8, color=NAVY, bold=True)
    separator = header_paragraph.add_run("  |  ENTERPRISE AI PRODUCT, OPERATIONS & ADOPTION LEADER")
    set_run_font(separator, size=7.5, color=MUTED)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    prefix = footer_paragraph.add_run("Page ")
    set_run_font(prefix, size=8, color=MUTED)
    add_page_field(footer_paragraph, "PAGE")
    middle = footer_paragraph.add_run(" of ")
    set_run_font(middle, size=8, color=MUTED)
    add_page_field(footer_paragraph, "NUMPAGES")


def add_header(doc):
    paragraph = doc.add_paragraph(style="Resume Name")
    run = paragraph.add_run("Alexander Aidun")
    set_run_font(run, size=26, color=NAVY, bold=True)

    title = doc.add_paragraph(style="Resume Title")
    title_run = title.add_run("Enterprise AI Product, Operations & Adoption Leader")
    set_run_font(title_run, size=12.5, color=BLUE, bold=True)

    contact = doc.add_paragraph(style="Resume Contact")
    set_run_font(contact.add_run("New York, NY  |  917-282-1668  |  "), size=9, color=MUTED)
    add_hyperlink(contact, "bobuel@gmail.com", "mailto:bobuel@gmail.com", color=MUTED)
    set_run_font(contact.add_run("  |  "), size=9, color=MUTED)
    add_hyperlink(contact, "linkedin.com/in/aaidun", "https://www.linkedin.com/in/aaidun/", color=MUTED)
    set_run_font(contact.add_run("  |  "), size=9, color=MUTED)
    add_hyperlink(contact, "aydoon.com", "https://aydoon.com", color=MUTED)
    set_run_font(contact.add_run("  |  "), size=9, color=MUTED)
    add_hyperlink(contact, "github.com/bobuel", "https://github.com/bobuel", color=MUTED)


def build_resume(output_path):
    doc = Document()
    doc.core_properties.title = "Alexander Aidun Resume"
    doc.core_properties.subject = "Enterprise AI Product, Operations and Adoption Leadership"
    doc.core_properties.author = "Alexander Aidun"
    doc.core_properties.keywords = (
        "enterprise AI, AI operations, AI product management, AI adoption, systems design"
    )
    doc.core_properties.comments = "Employer-facing resume"

    setup_styles(doc)
    setup_page_furniture(doc)
    bullet_num_id = add_custom_bullet_numbering(doc)
    add_header(doc)

    add_section_heading(doc, "Executive Profile")
    profile = doc.add_paragraph(style="Executive Profile")
    profile.add_run(
        "Enterprise AI product and adoption leader who designs the operating systems around AI: "
        "functional and cost operations, internal products, agentic workflows, learning, "
        "evaluation, and peer enablement. Currently helps operate a 1,500-employee AI "
        "environment at Automattic; previously led four AI initiatives at Dremio and scaled "
        "Dremio University to 3,200+ users with +78 NPS. Hands-on builder who translates "
        "emerging model capability into useful, repeatable work."
    )

    add_section_heading(doc, "Core Capabilities")
    capabilities = doc.add_paragraph(style="Compact Line")
    capabilities.add_run(
        "Enterprise AI operations | AI product strategy | Agentic workflows | Systems design | "
        "AI adoption and behavior change | Product discovery | Learning and enablement | "
        "Executive AI advisory | Cross-functional leadership | Hands-on prototyping"
    )
    tools_line = doc.add_paragraph(style="Compact Line")
    label = tools_line.add_run("Platforms and methods: ")
    set_run_font(label, size=9, color=NAVY, bold=True)
    tools_line.add_run(
        "ChatGPT, Claude, Codex, Claude Code, Cowork, MCP, Zapier, OpenAI, GitHub, Jira, and Pendo"
    )

    add_section_heading(doc, "Professional Experience")
    add_job(
        doc,
        company="Automattic",
        period="March 2026 - Present",
        role="AI Adoption Manager",
        location="New York",
        num_id=bullet_num_id,
        bullets=[
            "Help administer and operate ChatGPT (including Codex and Work) and Claude (including Cowork and Code) for 1,500 employees, with functional and cost responsibilities across the environment.",
            "Own product management for an internal AI Agent, AI Learning, LibreChat, and Slack-based agentic automations, connecting operations, workflows, guidance, and adoption.",
            "Design and deliver immersive enablement and publish 2-3 practical how-to articles per week to translate platform capability into repeatable work.",
            "Lead the AI Guides champion program and scope executive AI use cases as a personal AI product manager for C-suite leaders.",
        ],
    )
    add_job(
        doc,
        company="Dremio",
        period="January 2024 - March 2026",
        role="Senior AI Product Manager and Director, Education & Documentation",
        location="Remote",
        num_id=bullet_num_id,
        bullets=[
            "Led a four-initiative AI portfolio spanning an AI Agent, MCP Server, AI SQL Functions, and a Data Analyst chatbot; translated customer discovery into roadmap and workflow decisions.",
            "Partnered with Design and Engineering to prioritize use cases, product experience, and revenue potential across the portfolio.",
            "Owned the Documentation and Dremio University teams while driving cross-functional AI automation with Zapier, OpenAI, Jira, GitHub, and MCP.",
            "Separately, scaled Dremio University to 3,200+ users, 1,000+ badges, +78 NPS, and a 50% course completion rate within six months.",
        ],
    )
    add_job(
        doc,
        company="Braze",
        period="November 2022 - January 2024",
        role="Senior Manager, Partner Education",
        location="Remote",
        num_id=bullet_num_id,
        bullets=[
            "Led global partner education strategy, staffing, and programming across small, medium, and global strategic partners, including Accenture and Deloitte.",
            "Introduced generative AI and prompt engineering to replace live mock-session testing and accelerate content development.",
        ],
    )

    add_section_heading(doc, "Earlier Leadership Experience")
    add_job(
        doc,
        company="Arrikto",
        period="September 2021 - August 2022",
        role="Global Director, Education",
        location="Remote",
        num_id=bullet_num_id,
        bullets=[
            "Owned global customer and employee education strategy, staffing, and delivery for a data science and MLOps startup.",
            "Reached a 32% weekly course completion rate in six months, with more than 750 registrants and 1,400 enrollments.",
        ],
    )
    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)
    add_job(
        doc,
        company="WorkFusion",
        period="January 2021 - August 2021",
        role="Director, Product Enablement, Education & Documentation",
        location="New York, NY",
        num_id=bullet_num_id,
        bullets=[
            "Led distributed Documentation, Education, and Enablement functions across the United States, Belarus, Poland, and India, managing managers, writers, developers, and enablement staff.",
            "Connected product functionality to technical documentation, self-service hands-on learning, and coordinated enablement across a global operating model.",
        ],
    )
    add_job(
        doc,
        company="Qubole",
        period="December 2015 - December 2020",
        role="Global Director, Education Services & Technical Publications (2020); Global Director, Education Services (2015 - 2020)",
        location="New York, NY",
        num_id=bullet_num_id,
        bullets=[
            "Owned external customer education, employee technical onboarding, partner enablement and certification, and technical publications.",
            "Developed paid education services packages, including a largest engagement of $300K, and led a media producer, designer, and eight course developers.",
            "Pioneered in-product onboarding, education, help-center content, and telemetry with Pendo; partnered cross-functionally on employee development programs.",
        ],
    )

    add_section_heading(doc, "Additional Experience")
    additional = doc.add_paragraph(style="Compact Line")
    additional.add_run("MarketShare").bold = True
    additional.add_run(" - Manager, Data Strategy  |  ")
    additional.add_run("Data Meaning").bold = True
    additional.add_run(" - Technical Lead  |  ")
    additional.add_run("MicroStrategy").bold = True
    additional.add_run(" - Education Consultant; Technical Support Engineer")

    add_section_heading(doc, "Selected AI Builds")
    add_bullet(
        doc,
        "CertifyFast - Product concept for turning authoritative source material into traceable job-task analysis and exam-development artifacts for expert review.",
        bullet_num_id,
    )
    add_bullet(
        doc,
        "BloomGPT and Bloom Quiz Builder - Turned a 1,000+ use signal into a source-grounded assessment workflow with teacher checkpoints and structured output.",
        bullet_num_id,
    )
    add_bullet(
        doc,
        "Retrieval Guard - Experimental Python toolkit for retrieval regression testing and structurally similar near-miss filtering before context reaches an LLM or agent.",
        bullet_num_id,
    )

    add_section_heading(doc, "Education & Service")
    education = doc.add_paragraph(style="Compact Line")
    school = education.add_run("Cornell University, School of Engineering")
    set_run_font(school, size=9.2, color=NAVY, bold=True)
    education.add_run(" - Information Science, Systems and Technology")
    service = doc.add_paragraph(style="Compact Line")
    organization = service.add_run("Becket Chimney Corners YMCA")
    set_run_font(organization, size=9.2, color=NAVY, bold=True)
    service.add_run(" - Alumni Council / Ambassador, January 2017 - Present")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("docs/resume/alexander-aidun-resume.docx"),
    )
    args = parser.parse_args()
    build_resume(args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
